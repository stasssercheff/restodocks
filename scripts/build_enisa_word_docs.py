#!/usr/bin/env python3
"""Build ENISA technical brief Word documents from Markdown (ru/en/es)."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    para_buf: list[str] = []

    def flush_para() -> None:
        if not para_buf:
            return
        text = " ".join(para_buf).strip()
        para_buf.clear()
        if text:
            doc.add_paragraph(text)

    text = md_path.read_text(encoding="utf-8")
    for raw in text.splitlines():
        line = raw.rstrip("\n")
        stripped = line.strip()
        if stripped == "---":
            flush_para()
            continue
        if stripped.startswith("### "):
            flush_para()
            p = doc.add_heading(stripped[4:], level=3)
            p.runs[0].font.size = Pt(12)
            continue
        if stripped.startswith("## "):
            flush_para()
            p = doc.add_heading(stripped[3:], level=2)
            p.runs[0].font.size = Pt(13)
            continue
        if stripped.startswith("# "):
            flush_para()
            p = doc.add_heading(stripped[2:], level=1)
            p.runs[0].font.size = Pt(16)
            continue
        if stripped.startswith("- "):
            flush_para()
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue
        if not stripped:
            flush_para()
            continue
        para_buf.append(stripped)

    flush_para()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main() -> int:
    base = Path(__file__).resolve().parent / "enisa_docs"
    desktop = Path.home() / "Desktop"
    if not desktop.is_dir():
        desktop = Path.home() / "Escritorio"  # Spanish macOS folder name
    pairs = [
        (base / "ru.md", desktop / "Restodocks_ENISA_Technical_ru.docx"),
        (base / "en.md", desktop / "Restodocks_ENISA_Technical_en.docx"),
        (base / "es.md", desktop / "Restodocks_ENISA_Technical_es.docx"),
    ]
    for md, out in pairs:
        if not md.is_file():
            print(f"Missing {md}", file=sys.stderr)
            return 1
        md_to_docx(md, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
